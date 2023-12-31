##-------------------------------------------------------------------------------------##
## Include
##-------------------------------------------------------------------------------------## 
#import pyaudio # Only on device, not on client
#import simpleaudio # Only on device, not on client
#import wave # Only on device, not on client
#from pygame import mixer # Only on device, not on client
from array import array

from gtts import gTTS
from openai import OpenAI

import os 
from pathlib import Path

import time
import datetime
import base64

import math
import numpy as np
import matplotlib.pyplot as plt
from docx import Document

import streamlit as st
#from audio_recorder_streamlit import audio_recorder
from streamlit_mic_recorder import mic_recorder
from streamlit.logger import get_logger
from typing import Any
import requests
import pandas as pd
import zipfile

##-------------------------------------------------------------------------------------##
## OpenAI
##-------------------------------------------------------------------------------------## 
localSecrets = Path("./localSecrets.txt")
if localSecrets.is_file():
    f = open(localSecrets, "r")
    localKey = str(f.read())
    os.environ['OPENAI_API_KEY'] = localKey

# ##-------------------------------------------------------------------------------------##
# ## playWav (only on device not client)
# ##-------------------------------------------------------------------------------------##   
# def playWav(filename):
#     wave_obj = simpleaudio.WaveObject.from_wave_file(filename)
#     play_obj = wave_obj.play()
#     play_obj.wait_done()
    
# ##-------------------------------------------------------------------------------------##
# ## playAudio (only on device not client)
# ##-------------------------------------------------------------------------------------##   
# def playAudio(filename, background=False):

#     mixer.init()
#     mixer.music.load(filename)
#     mixer.music.play()
#     while not background:
#         if mixer.music.get_busy() == False:
#             #mixer.quit()
#             break

# ##-------------------------------------------------------------------------------------##
# ## stopAudio (only on device not client)
# ##-------------------------------------------------------------------------------------##   
# def stopAudio( ):
#     if mixer.get_init():
#         mixer.music.stop()
#         mixer.quit()

##-------------------------------------------------------------------------------------##
## playAudioEmbedded
##-------------------------------------------------------------------------------------##   
# TODO Change playbackRate
def playAudioEmbedded(filename, background=False):

    try:
        with open(filename, "rb") as f:
            data = f.read()
            b64 = base64.b64encode(data).decode()
            # <audio controls autoplay="true"> in case you want to show controls
            md = f"""
                <audio id="audio" controls autoplay="true" style="width: 100%;">
                <source src="data:audio/mp3;base64,{b64}" type="audio/mp3">
                </audio>
                """
            st.markdown(
                md,
                unsafe_allow_html=True,
            )
    except FileNotFoundError:
        print("File not found. Check the path variable and filename")
        return    
     
# ##-------------------------------------------------------------------------------------##
# ## recordAudio (only on device not client)
# ##-------------------------------------------------------------------------------------##     
# def recordAudio(outputFile, NOISE=100, SILENCE=2):
#     # Defining audio variables
#     CHUNK = 1024
#     FORMAT = pyaudio.paInt16
#     CHANNELS = 1
#     RATE = 44100
      
#     # Calling pyadio module and starting recording 
#     p = pyaudio.PyAudio()
#     stream = p.open(format=FORMAT,
#                 channels=CHANNELS, 
#                 rate=RATE, 
#                 input=True,
#                 frames_per_buffer=CHUNK)
 
#     stream.start_stream()
 
#     # Recording data until under threshold
#     frames=[]
#     isilent = 0
#     isilentmax = int(RATE / CHUNK * SILENCE)
#     while True:
#         data=stream.read(CHUNK)
#         data_chunk = array('h',data)
#         volume = max(data_chunk)
#         frames.append(data)
#         if volume < NOISE:
#             isilent = isilent + 1
#         else:
#             isilent = 0
#         if isilent > isilentmax: 
#             break
#     done = 1
#     if ( len(frames)-1 == isilentmax ): done = -1
            
#     # Stopping recording   
#     stream.stop_stream()
#     stream.close()
#     p.terminate()
    
#     # Saving file with wave module    
#     wf = wave.open(outputFile, 'wb')
#     wf.setnchannels(CHANNELS)
#     wf.setsampwidth(p.get_sample_size(FORMAT))
#     wf.setframerate(RATE)
#     wf.writeframes(b''.join(frames))
#     wf.close()
#     return done

##-------------------------------------------------------------------------------------##
## speechToTextGoogle
##-------------------------------------------------------------------------------------##
import speech_recognition as sr
def speechToTextGoogle(filename, language="en"):
    
    transcript = ""
    r = sr.Recognizer()
    with sr.Microphone() as source:
        # Adjust the energy threshold based on the surrounding noise level
        r.adjust_for_ambient_noise(source, duration=0.2)
                
        # Listens for the user's input
        question = r.listen(source)
                
        # Using google to recognize audio
        transcript = r.recognize_google(question, language=language)
        return transcript

##-------------------------------------------------------------------------------------##
## speechToTextOpenAI
##-------------------------------------------------------------------------------------##
def speechToTextOpenAI(filename, language="en"):
    client = OpenAI()
    audio_file= open(filename, "rb")
    transcript = client.audio.transcriptions.create(
        model="whisper-1", 
        language=language,
        file=audio_file,
        response_format="text"
    )
    
    # Controls
    if transcript.find("Amara.org") > 0: transcript = " "
    if transcript.find("prossimo episodio") > 0: transcript = " "

    return transcript

##-------------------------------------------------------------------------------------##
## chatBot
##-------------------------------------------------------------------------------------##
def chatBot(input, language = "en"):

    prompt = "You are an helpful virtual assistant. Answer with max 50 words"
    if ( language == "it" ): prompt = "Sei un assistente virtuale. Rispondi al massimo con 50 parole"
    client = OpenAI()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
        {"role": "system", "content": prompt },
        #{"role": "system", "content": "You are the best children story author in the world. Write a story about this topic"},
        {"role": "user", "content": input}
        ]
    )
    return(completion.choices[0].message.content)

##-------------------------------------------------------------------------------------##
## summarizeBot
##-------------------------------------------------------------------------------------##
def summarizeBot(input, language = "en"):

    prompt = "Summarize the provided text"
    if ( language == "it" ): prompt = "Riassumi il testo fornito"
    client = OpenAI()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
        {"role": "system", "content": prompt},
        {"role": "user", "content": input}
        ]
    )
    return(completion.choices[0].message.content)

##-------------------------------------------------------------------------------------##
## sentimentAnalysis
##-------------------------------------------------------------------------------------##
def sentimentAnalysis(transcript):
    content = f"What emotion is the following text expressing?\n{transcript}"
    client = OpenAI()
    completion = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
        {"role": "system", "content": "You are a helpful assistant. Answer with one word choosing between Positive, Neutral, Negative"},
        {"role": "user", "content": content}
        ]
    )
    return(completion.choices[0].message.content)

##-------------------------------------------------------------------------------------##
## textToSpeechGoogle
##-------------------------------------------------------------------------------------##
def textToSpeechGoogle(text, filename, language="en"):

    # Check audio file is not open
    while True:
        try:
            myfile = open(filename, "wb")
            break                             
        except IOError:
            #stopAudio( )
            return

    speech = gTTS(text, lang=language)
    speech.save(filename)

##-------------------------------------------------------------------------------------##
## textToSpeechOpenAI
##-------------------------------------------------------------------------------------##
def textToSpeechOpenAI(text, filename, language="alloy"):
    
    # Check audio file is not open
    while True:
        try:
            myfile = open(filename, "wb")
            break                             
        except IOError:
            #stopAudio( ) 
            return 

    speech_file_path = Path(__file__).parent / filename
    client = OpenAI()
    response = client.audio.speech.create(
        model="tts-1",
        voice=language, # alloy, onyx, fable
        input=text
    )
    response.stream_to_file(speech_file_path)

##-------------------------------------------------------------------------------------##
## verticalSpaceSideBar
##-------------------------------------------------------------------------------------##
def verticalSpaceSideBar():
 
    # Extend vertical spacing of sidebar
    st.sidebar.markdown("""
        <style>
        [data-testid='stSidebarNav'] > ul {
            min-height: 45vh;
        } 
        </style>
        """, unsafe_allow_html=True)

##-------------------------------------------------------------------------------------##
## createReport
##-------------------------------------------------------------------------------------##
def createReport(language = "en"):
 
    # Read 
    with open('Questions.txt', 'r') as f:
        Questions=f.read()
    n = 0
    Sentiments = [ int(0) ]*1000
    with open('Sentiments.txt', 'r') as f:
        for line in f:
            Sentiments[n] = int(line)
            n += 1
    Sentiments = Sentiments[0:n]

    # Post processing
    Experience = [ int(0) ]*n
    xAxis = [ int(0) ]*n
    Colors = [ 'white' ]*n
    Experience[0] = 0 #Sentiments[0]
    for k in range(1, n):
        Experience[k] = Experience[k-1] + Sentiments[k-1]
        xAxis[k] = k
        if ( Sentiments[k]>0 ): Colors[k] = "green"
        if ( Sentiments[k]==0 ): Colors[k] = "yellow"
        if ( Sentiments[k]<0 ): Colors[k] = "red"

    # Creating the bar plot
    fig = plt.figure()
    yticks = range(math.floor(min(Experience)-1), math.ceil(max(Experience)+1))
    plt.yticks(yticks)
    plt.grid(visible=True, axis='y', linestyle='--')
    #plt.ylim(-1.1, 1.1)
    plt.bar(xAxis, Sentiments, bottom = Experience, color = Colors)
    plt.xlabel("Interactions")
    plt.ylabel("Sentiment evolution")
    plt.savefig('Experience.png')

    # Summarize
    summary = summarizeBot(Questions, language)
    # document = Document()
    # document.add_heading("Sintesi discussione voicebot " + str(datetime.date.today()), level=1)
    # p = document.add_paragraph(summary)
    # p = document.add_picture('Experience.png')
    # document.save('Summary.docx')

    return summary

##-------------------------------------------------------------------------------------##
## createStoriesMp3
##-------------------------------------------------------------------------------------##
def createStoriesMp3():

    # Initialize progress bar
    myprogress = st.sidebar.progress(0, "Creating audio book...")
    
    stories = os.listdir("./stories")
    outputFile = "myStories.mp3"
    tempFile = "temp.mp3"
    voice = "alloy"
    N = len(stories)
    mf = open(outputFile, "wb")
    mf.close()
    for k in range(0, N):
        filepath = "./stories/" + str(stories[k])
        fid = open(filepath, 'r', encoding='latin-1')
        text = fid.read()
        temp = text.split("\n",3)
        fid.close()
        language = temp[2]
        body = temp[3]

        if ( ( voice == "alloy" or voice == "onyx" or voice == "fable" ) and ( len(body) < 4096 ) ):
            textToSpeechOpenAI(body, tempFile, language=voice)
        else:
            language_code = "en"
            if language == "Italian": language_code = "it"
            textToSpeechGoogle(body, tempFile, language_code)

        with open(outputFile, "ab") as mf, open(tempFile, "rb") as tf:
            mf.write(tf.read())

        myprogress.progress(float((k+1)/N),"Creating audio book...")

##-------------------------------------------------------------------------------------##
## createStoriesWord
##-------------------------------------------------------------------------------------##
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Cm, Pt
def createStoriesWord():

    # Initialize progress bar
    myprogress = st.sidebar.progress(0, "Creating book...")
    
    document = Document()
    font = document.styles['Normal'].font
    font.size = Pt(16)
    stories = os.listdir("./stories")
    N = len(stories)
    for k in range(0, N):
        filepath = "./stories/" + str(stories[k])
        fid = open(filepath, 'r', encoding='latin-1')
        text = fid.read()
        temp = text.split("\n",4)
        fid.close()
        title = temp[3]
        body = temp[4]

        document.add_heading(title, level=0)
        paragraphs = body.splitlines()
        for j in range(0, len(paragraphs)):
            if ( len(paragraphs[j]) > 1 ):
               paragraph = document.add_paragraph(paragraphs[j])
               paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
        picturePath = Path("./images/" + stories[k] + ".png")
        if picturePath.is_file():   
            picture = document.add_picture("./images/" + stories[k] + ".png", width=Inches(6))
            last_paragraph = document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_page_break()
    
        myprogress.progress(float((k+1)/N),"Creating book...")
    
    document.save('myStories.docx')

##-------------------------------------------------------------------------------------##
## createStoriesPdf
##-------------------------------------------------------------------------------------##
from pdfme import build_pdf
def createStoriesPdf():

    # Initialize progress bar
    myprogress = st.sidebar.progress(0, "Creating book...")
    
    document = {
    "style": {
        "margin_bottom": 15, "text_align": "j",
        "page_size": "letter", "margin": [60, 50]
    },
    "formats": {
        "url": {"c": "blue", "u": 1},
        "title": {"b": 1, "s": 30}
    },
    "running_sections": {
        "footer": {
            "x": "left", "y": 740, "height": "bottom", "style": {"text_align": "c"},
            "content": [{".": ["Page ", {"var": "$page"}]}]
        },
    },
    "sections": [ 
        {"style": { "s":15, "f":"Times"},
        "running_sections": ["footer"],
        "content": [
            {".": "My book of bedtime stories", "style": "title"},
            ["Published on " + str(datetime.date.today())],
            ]
    }],
    }

    stories = os.listdir("./stories")
    N = len(stories)
    for k in range(0, N):
        filepath = "./stories/" + str(stories[k])
        fid = open(filepath, 'r', encoding='latin-1')
        text = fid.read()
        temp = text.split("\n",4)
        fid.close()
        title = temp[3]
        body = temp[4]

        document["sections"].append({
            "style": { "s":14, "f":"Times"},
            "running_sections": ["footer"],
            "content": [
                {".": title, "style": "title"},
                [body],
                {"image": "./images/" + stories[k] + ".png"},
                ]
        })
  
        myprogress.progress(float((k+1)/N),"Creating book...")
    
    with open('myStories.pdf', 'wb') as f:
        build_pdf(document, f)

##-------------------------------------------------------------------------------------##
## displayPDF
##-------------------------------------------------------------------------------------##
def displayPDF(filename):
    
    # Opening file from file path
    with open(filename, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    # Embedding PDF in HTML
    pdf_display = F'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="940" type="application/pdf"></iframe>'
    #pdf_display = F'<iframe src="https://thomasmorestudies.org/wp-content/uploads/2020/09/Richard.pdf" width="100%" height="940" type="application/pdf"></iframe>'
    #pdf_display = F'<iframe src="./stories/myStories.pdf" width="100%" height="940" type="application/pdf"> </iframe>'
    # Displaying File
    st.markdown(pdf_display, unsafe_allow_html=True)

# def displayPDF(filename):
#     # Opening file from file path
#     with open(filename, "rb") as f:
#         base64_pdf = base64.b64encode(f.read()).decode('utf-8')
#     # Embedding PDF in HTML
#     pdf_display =  f"""<embed
#     class="pdfobject"
#     type="application/pdf"
#     title="Embedded PDF"
#     src="data:application/pdf;base64,{base64_pdf}"
#     style="overflow: auto; width: 100%; height: 940;">"""
#     # Displaying File
#     st.markdown(pdf_display, unsafe_allow_html=True)

##-------------------------------------------------------------------------------------##
## zipFolder
##-------------------------------------------------------------------------------------##
def zipFolder( dirpath, filename ):
    zf = zipfile.ZipFile(filename, "w")
    for dirname, subdirs, files in os.walk(dirpath):
        zf.write(dirname)
        for ff in files:
            zf.write(os.path.join(dirname, ff))
    zf.close()
    return

##-------------------------------------------------------------------------------------##
## unzipFolder
##-------------------------------------------------------------------------------------##
def unzipFolder( filename, dirpath ):
    with zipfile.ZipFile(filename,"r") as zf:
        zf.extractall(dirpath)
    return  

##-------------------------------------------------------------------------------------##
## chatVoiceBot
##-------------------------------------------------------------------------------------##
def chatVoiceBot(language = "IT"):

    # Central alignment of the spinner
    st.markdown("""
    <style>
    div.stSpinner > div {
        text-align:center;
        align-items: center;
        justify-content: center;
    }
    </style>""", unsafe_allow_html=True)

    # Initiatlize environment
    if 'k' not in st.session_state:
        st.session_state['k'] = str(0)
        fq = open('Questions.txt', 'w'); fq.close()
        fs = open('Sentiments.txt', 'w'); fs.close()
    inputFile = "input.wav" 
    outputFile = "output.mp3"

    # Initialize session
    withAudio = st.sidebar.toggle('Audio')
    withMic = st.sidebar.toggle('Mic')
    if "messages" not in st.session_state:
        st.session_state.messages = []
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    if ( withAudio and withMic ): playAudioEmbedded(outputFile)

    # Chat prompt
    prompt = ""
    prompt = st.chat_input("Write here...")

    # Audio prompt
    volume = -1
    question = ""
    if ( ( not prompt ) and withMic ): 
        
        audio = mic_recorder(start_prompt="Start recording...", stop_prompt="Stop recording...", key='recorder', just_once=True, use_container_width=True)
        if audio: 
            audio_bytes = audio['bytes']
            volume = len(audio_bytes)
        if ( volume > 0 ):
            with open(inputFile, mode='wb') as f:
                f.write(audio_bytes)
            question=speechToTextOpenAI(inputFile, language)
            
        # #Version listening in the background (valid locally)     
        # with st.spinner('\n Speak up \n'):
        #     volume = recordAudio(inputFile)
        #     if ( volume > 0 ): question=speechToTextOpenAI(inputFile, language)
        #     else: 
        #         time.sleep(1) 
        #         st.rerun()
    if ( len(question) > 0 ): prompt = question 

    if prompt:

        # Chat message question
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        # Commands
        question = prompt
        answer = ""
        if ( question == "/summary" ):
            answer = createReport(language)
            answer = "Below a summary of the conversation \n \n \n" + answer
        elif ( question == "/sentiment" ):
            answer = "Below a summary of the sentiment of the conversation \n "
            createReport(language)
        else:
            # Chatbot
            if ( len(question) > 0 ): answer=chatBot(question, language)
    
        # Chat message answer
        st.session_state.messages.append({"role": "assistant", "content": answer})
        with st.chat_message("assistant"):
            st.markdown(answer)
        if ( question == "/sentiment" ): st.image('Experience.png')
        
        # Play audio
        if ( withAudio and len(answer) > 0 ): textToSpeechGoogle(answer, outputFile, language)
        #if ( withAudio and len(answer) > 0 ): textToSpeechOpenAI(answer, outputFile, "onyx")
        #if ( withAudio and len(answer) > 0 and ( not withMic ) ): playAudio(outputFile)
        if ( withAudio and len(answer) > 0 and ( not withMic ) ): playAudioEmbedded(outputFile)

        # Sentiment analysis
        sentiment = ""
        sentiment = sentimentAnalysis(question)
        sentimentId = 0; 
        if sentiment == "Positive": sentimentId = 1
        if sentiment == "Negative": sentimentId = -1

        # Store results
        if ( 'k' in st.session_state ):
            fq = open('Questions.txt', 'a'); fq.write(question + '\n'); fq.close()
            fs = open('Sentiments.txt', 'a'); fs.write(str(sentimentId) + '\n'); fs.close()
            st.session_state['k'] = str(int(st.session_state['k']) + 1)

    # Refresh if mic used 
    if ( withMic and volume > 0 ): 
        st.rerun()
